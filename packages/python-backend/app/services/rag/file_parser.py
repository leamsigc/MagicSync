import io
import csv
import json
import logging
from dataclasses import dataclass, field
from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "text/html",
    "text/csv",
    "application/json",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".html": "text/html",
    ".htm": "text/html",
    ".csv": "text/csv",
    ".json": "application/json",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@dataclass
class ParsedPage:
    """A single page/section of a parsed document."""
    content: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Full parsed document with structured pages."""
    text: str
    pages: list[ParsedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


def extract_text(content: bytes, mime_type: str) -> str:
    """Extract plain text from file bytes based on MIME type.

    Args:
        content: Raw file bytes
        mime_type: MIME type of the file

    Returns:
        Extracted plain text

    Raises:
        ValueError: If the MIME type is not supported
    """
    result = extract_structured(content, mime_type)
    return result.text


def extract_structured(content: bytes, mime_type: str) -> ParsedDocument:
    """Extract structured text with page/section metadata from file bytes.

    Args:
        content: Raw file bytes
        mime_type: MIME type of the file

    Returns:
        ParsedDocument with text and per-page metadata
    """
    if mime_type not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {mime_type}")

    parsers = {
        "application/pdf": _extract_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
        "text/html": _extract_html,
        "text/markdown": _extract_markdown,
        "text/plain": _extract_plain,
        "text/csv": _extract_csv,
        "application/json": _extract_json,
    }

    parser = parsers.get(mime_type)
    if not parser:
        raise ValueError(f"No parser for: {mime_type}")

    return parser(content)


def _extract_pdf(content: bytes) -> ParsedDocument:
    """Extract text from PDF using pypdf with page-level metadata."""
    reader = PdfReader(io.BytesIO(content))
    pages = []
    all_text_parts = []

    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text and text.strip():
                cleaned = _clean_whitespace(text)
                pages.append(ParsedPage(
                    content=cleaned,
                    page_number=i + 1,
                    metadata={"source": "pdf", "total_pages": len(reader.pages)},
                ))
                all_text_parts.append(cleaned)
        except Exception as e:
            logger.warning(f"Failed to extract PDF page {i + 1}: {e}")
            continue

    metadata = {"total_pages": len(reader.pages), "extracted_pages": len(pages)}
    if reader.metadata:
        if reader.metadata.title:
            metadata["pdf_title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["pdf_author"] = reader.metadata.author

    return ParsedDocument(
        text="\n\n".join(all_text_parts),
        pages=pages,
        metadata=metadata,
    )


def _extract_docx(content: bytes) -> ParsedDocument:
    """Extract text from DOCX with heading-based section metadata."""
    doc = DocxDocument(io.BytesIO(content))
    pages = []
    all_text_parts = []
    current_section = None
    current_section_text = []
    page_number = 1

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        # Detect headings as section boundaries
        is_heading = para.style.name.startswith("Heading") if para.style else False

        if is_heading:
            # Save previous section
            if current_section_text:
                section_content = "\n\n".join(current_section_text)
                pages.append(ParsedPage(
                    content=section_content,
                    page_number=page_number,
                    section_title=current_section,
                    metadata={"source": "docx", "style": "body"},
                ))
                all_text_parts.append(section_content)
                page_number += 1

            current_section = para.text.strip()
            current_section_text = [current_section]  # Include heading in text flow
        else:
            current_section_text.append(para.text.strip())

    # Final section
    if current_section_text:
        section_content = "\n\n".join(current_section_text)
        pages.append(ParsedPage(
            content=section_content,
            page_number=page_number,
            section_title=current_section,
            metadata={"source": "docx", "style": "body"},
        ))
        all_text_parts.append(section_content)

    # If no headings found, treat whole doc as one page
    if not pages:
        full_text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        pages.append(ParsedPage(
            content=full_text,
            page_number=1,
            metadata={"source": "docx"},
        ))
        all_text_parts.append(full_text)

    return ParsedDocument(
        text="\n\n".join(all_text_parts),
        pages=pages,
        metadata={"paragraphs": len(doc.paragraphs), "sections": len(pages)},
    )


def _extract_html(content: bytes) -> ParsedDocument:
    """Extract text from HTML with heading-based sections."""
    text = content.decode("utf-8", errors="replace")
    soup = BeautifulSoup(text, "lxml")

    # Remove noise elements
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "noscript"]):
        tag.decompose()

    pages = []
    all_text_parts = []
    page_number = 1

    # Split by heading elements for section-based chunking
    current_section = None
    current_section_text = []

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "pre", "blockquote", "div"]):
        if element.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            # Save previous section
            if current_section_text:
                section_content = "\n\n".join(current_section_text)
                if section_content.strip():
                    pages.append(ParsedPage(
                        content=section_content,
                        page_number=page_number,
                        section_title=current_section,
                        metadata={"source": "html"},
                    ))
                    all_text_parts.append(section_content)
                    page_number += 1

            current_section = element.get_text(strip=True)
            current_section_text = [current_section]  # Include heading in text flow
        else:
            line = element.get_text(strip=True)
            if line:
                current_section_text.append(line)

    # Final section
    if current_section_text:
        section_content = "\n\n".join(current_section_text)
        if section_content.strip():
            pages.append(ParsedPage(
                content=section_content,
                page_number=page_number,
                section_title=current_section,
                metadata={"source": "html"},
            ))
            all_text_parts.append(section_content)

    # Fallback: if no headings found, extract all text
    if not pages:
        lines = soup.get_text(separator="\n").split("\n")
        cleaned = [line.strip() for line in lines if line.strip()]
        full_text = "\n\n".join(cleaned)
        pages.append(ParsedPage(
            content=full_text,
            page_number=1,
            metadata={"source": "html"},
        ))
        all_text_parts.append(full_text)

    return ParsedDocument(
        text="\n\n".join(all_text_parts),
        pages=pages,
        metadata={"sections": len(pages)},
    )


def _extract_markdown(content: bytes) -> ParsedDocument:
    """Extract text from Markdown preserving heading structure."""
    text = content.decode("utf-8", errors="replace")

    # Strip YAML frontmatter if present
    frontmatter = {}
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            # Parse basic YAML frontmatter
            for line in parts[1].strip().split("\n"):
                if ":" in line:
                    key, _, value = line.partition(":")
                    frontmatter[key.strip()] = value.strip()
            text = parts[2].strip()

    pages = []
    all_text_parts = []
    page_number = 1
    current_heading = None
    current_section_lines = []

    for line in text.split("\n"):
        heading_match = line.lstrip()
        if heading_match.startswith("#"):
            # Save previous section
            if current_section_lines:
                section_content = "\n".join(current_section_lines).strip()
                if section_content:
                    pages.append(ParsedPage(
                        content=section_content,
                        page_number=page_number,
                        section_title=current_heading,
                        metadata={"source": "markdown", **frontmatter},
                    ))
                    all_text_parts.append(section_content)
                    page_number += 1

            # Extract heading text (strip # prefix)
            current_heading = heading_match.lstrip("#").strip()
            current_section_lines = [current_heading]  # Include heading in text flow
        else:
            current_section_lines.append(line)

    # Final section
    if current_section_lines:
        section_content = "\n".join(current_section_lines).strip()
        if section_content:
            pages.append(ParsedPage(
                content=section_content,
                page_number=page_number,
                section_title=current_heading,
                metadata={"source": "markdown", **frontmatter},
            ))
            all_text_parts.append(section_content)

    # If no headings, convert full text via markdown parser
    if not pages:
        html = markdown.markdown(text, extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html, "lxml")
        cleaned = _clean_whitespace(soup.get_text(separator="\n"))
        pages.append(ParsedPage(
            content=cleaned,
            page_number=1,
            metadata={"source": "markdown", **frontmatter},
        ))
        all_text_parts.append(cleaned)

    return ParsedDocument(
        text="\n\n".join(all_text_parts),
        pages=pages,
        metadata={"sections": len(pages), **frontmatter},
    )


def _extract_plain(content: bytes) -> ParsedDocument:
    """Extract plain text."""
    text = content.decode("utf-8", errors="replace")
    cleaned = _clean_whitespace(text)
    return ParsedDocument(
        text=cleaned,
        pages=[ParsedPage(content=cleaned, page_number=1, metadata={"source": "plain"})],
        metadata={"source": "plain"},
    )


def _extract_csv(content: bytes) -> ParsedDocument:
    """Extract text from CSV with row-based structure."""
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ParsedDocument(text="", pages=[], metadata={"rows": 0, "columns": 0})

    headers = rows[0]
    data_rows = rows[1:]

    pages = []
    all_text_parts = []

    # Convert each row to a natural language description
    for i, row in enumerate(data_rows):
        if not any(cell.strip() for cell in row):
            continue

        parts = []
        for header, value in zip(headers, row):
            if value.strip():
                parts.append(f"{header.strip()}: {value.strip()}")

        if parts:
            row_text = "; ".join(parts)
            pages.append(ParsedPage(
                content=row_text,
                page_number=i + 1,
                metadata={"source": "csv", "row_index": i},
            ))
            all_text_parts.append(row_text)

    # Also create a summary table as text
    summary_lines = [", ".join(headers)]
    for row in data_rows[:100]:  # Limit to first 100 rows for summary
        if any(cell.strip() for cell in row):
            summary_lines.append(", ".join(row))

    summary = "\n".join(summary_lines)

    return ParsedDocument(
        text="\n\n".join(all_text_parts) if all_text_parts else summary,
        pages=pages if pages else [ParsedPage(content=summary, page_number=1, metadata={"source": "csv"})],
        metadata={"rows": len(data_rows), "columns": len(headers), "headers": headers},
    )


def _extract_json(content: bytes) -> ParsedDocument:
    """Extract text from JSON with structured traversal."""
    text = content.decode("utf-8", errors="replace")

    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON: {e}")

    pages = []
    all_text_parts = []

    def flatten_json(obj, prefix: str = "") -> list[tuple[str, str]]:
        """Flatten JSON into key-value text pairs."""
        items = []
        if isinstance(obj, dict):
            for key, value in obj.items():
                full_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    items.extend(flatten_json(value, full_key))
                else:
                    items.append((full_key, str(value)))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                full_key = f"{prefix}[{i}]" if prefix else f"[{i}]"
                if isinstance(item, (dict, list)):
                    items.extend(flatten_json(item, full_key))
                else:
                    items.append((full_key, str(item)))
        else:
            items.append((prefix, str(obj)))
        return items

    if isinstance(data, list):
        # Array of objects — each object is a page
        for i, item in enumerate(data):
            if isinstance(item, dict):
                parts = [f"{k}: {v}" for k, v in item.items() if not isinstance(v, (dict, list))]
                text_content = "; ".join(parts)
            else:
                text_content = str(item)

            if text_content.strip():
                pages.append(ParsedPage(
                    content=text_content,
                    page_number=i + 1,
                    metadata={"source": "json", "index": i},
                ))
                all_text_parts.append(text_content)
    elif isinstance(data, dict):
        # Single object — flatten into sections
        flattened = flatten_json(data)
        page_content = "\n".join(f"{k}: {v}" for k, v in flattened)
        pages.append(ParsedPage(
            content=page_content,
            page_number=1,
            metadata={"source": "json"},
        ))
        all_text_parts.append(page_content)
    else:
        text_content = str(data)
        pages.append(ParsedPage(
            content=text_content,
            page_number=1,
            metadata={"source": "json"},
        ))
        all_text_parts.append(text_content)

    return ParsedDocument(
        text="\n\n".join(all_text_parts),
        pages=pages,
        metadata={"source": "json", "type": type(data).__name__},
    )


def _clean_whitespace(text: str) -> str:
    """Collapse whitespace and filter empty lines."""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            cleaned.append(stripped)
    return "\n\n".join(cleaned)
