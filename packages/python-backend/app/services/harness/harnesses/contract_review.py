import logging
import json
from typing import Any, Dict

logger = logging.getLogger(__name__)


async def extract_document(phase_input: Any, context: Dict) -> Dict:
    """Phase 1: Document Intake - Extract text from uploaded DOCX/PDF."""
    document_path = phase_input.get("path") or phase_input.get("document_id")
    
    if not document_path:
        return {"error": "No document path provided"}
    
    try:
        text_content = ""
        mime_type = ""
        
        if document_path.endswith(".pdf"):
            try:
                from pypdf import PdfReader
                reader = PdfReader(document_path)
                text_content = "\n".join([page.extract_text() for page in reader.pages])
                mime_type = "application/pdf"
            except Exception as e:
                return {"error": f"Failed to extract PDF: {e}"}
        
        elif document_path.endswith((".docx", ".doc")):
            try:
                from docx import Document
                doc = Document(document_path)
                text_content = "\n".join([para.text for para in doc.paragraphs])
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except Exception as e:
                return {"error": f"Failed to extract DOCX: {e}"}
        
        else:
            return {"error": f"Unsupported file type: {document_path}"}
        
        return {
            "extracted_text": text_content[:5000],
            "mime_type": mime_type,
            "document_path": document_path,
            "char_count": len(text_content)
        }
    
    except Exception as e:
        logger.error(f"Document extraction failed: {e}")
        return {"error": str(e)}


async def classify_contract(phase_input: Any, context: Dict) -> Dict:
    """Phase 2: Classification - Classify contract type using LLM."""
    from app.services.llm import llm_service
    
    text = phase_input.get("extracted_text", "")[:3000]
    
    prompt = f"""Classify this contract. Output as JSON:
{{
    "contract_type": "NDA|Employment|Service|Lease|Purchase|Other",
    "confidence": 0.0-1.0,
    "key_parties": ["party1", "party2"],
    "summary": "2-3 sentence summary"
}}

Contract text:
{text}"""
    
    response = await llm_service.chat([
        {"role": "system", "content": "You classify contracts. Output valid JSON only."},
        {"role": "user", "content": prompt}
    ], json_mode=True)
    
    try:
        classification = json.loads(response)
        return {"classification": classification}
    except Exception:
        return {"classification": {"contract_type": "Unknown", "confidence": 0.0}}


async def extract_clauses(phase_input: Any, context: Dict) -> Dict:
    """Phase 5: Clause Extraction - Parse contract into clauses."""
    text = phase_input.get("extracted_text", phase_input.get("text", ""))[:5000]
    
    import re
    
    clause_patterns = [
        r'(?i)(?:^|\n)(?:ARTICLE|Clause|Section)\s*(\d+|[a-z])\.?\s*([^\n]+)',
        r'(?i)(?<=^|\n\n)([A-Z][A-Z\s]+?)(?=\n\n)',
    ]
    
    clauses = []
    
    for pattern in clause_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            clauses.append({
                "identifier": match.group(1) if match.lastindex >= 1 else "",
                "title": match.group(2) if match.lastindex >= 2 else match.group(1),
                "text": match.group(0)[:500]
            })
    
    if not clauses:
        sections = text.split("\n\n")
        clauses = [{"identifier": str(i), "title": "Section", "text": s[:500]} 
                   for i, s in enumerate(sections[:20]) if s.strip()]
    
    return {"clauses": clauses, "count": len(clauses)}


async def generate_summary(phase_input: Any, context: Dict) -> Dict:
    """Phase 8: Executive Summary - Generate summary + DOCX report."""
    from app.services.llm import llm_service
    
    classification = phase_input.get("classification", {})
    clauses = phase_input.get("clauses", [])
    risk_analysis = phase_input.get("risk_analysis", {})
    
    prompt = f"""Generate an executive summary of this contract review.

Contract Type: {classification.get('contract_type', 'Unknown')}
Key Parties: {classification.get('key_parties', [])}
Risk Level: {risk_analysis.get('overall_risk', 'Unknown')}
Clauses Analyzed: {len(clauses)}

Provide a 3-paragraph executive summary covering:
1. Overview and purpose
2. Key risks and concerns
3. Recommendations

Then output final JSON:
{{"summary": "...", "recommendations": ["...", "..."], "risk_level": "low|medium|high"}}"""

    response = await llm_service.chat([
        {"role": "system", "content": "You generate contract review summaries."},
        {"role": "user", "content": prompt}
    ], json_mode=True)
    
    try:
        result = json.loads(response)
        
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Contract Review Report', 0)
            doc.add_paragraph(f"Contract Type: {classification.get('contract_type', 'Unknown')}")
            doc.add_paragraph(f"Risk Level: {result.get('risk_level', 'Unknown')}")
            doc.add_heading('Executive Summary', 1)
            doc.add_paragraph(result.get('summary', ''))
            doc.add_heading('Recommendations', 1)
            for rec in result.get('recommendations', []):
                doc.add_paragraph(rec, style='List Bullet')
            
            report_path = f"/tmp/contract_review_{context.get('run_id', 'report')}.docx"
            doc.save(report_path)
            result['report_path'] = report_path
        except Exception as e:
            logger.warning(f"DOCX generation failed: {e}")
        
        return result
    except Exception:
        return {"summary": "Summary generation failed", "recommendations": [], "risk_level": "unknown"}


CONTRACT_REVIEW_HANDLERS = {
    "extract_document": extract_document,
    "classify_contract": classify_contract,
    "extract_clauses": extract_clauses,
    "generate_summary": generate_summary,
}
